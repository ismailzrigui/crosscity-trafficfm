from __future__ import annotations

import csv
import os
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEX = ROOT / "paper" / "main.tex"
BIB = ROOT / "paper" / "references.bib"
OUT = Path(os.environ.get("DOCX_OUT", ROOT / "results" / "manuscript" / "CrossCity_TrafficFM_article.docx"))

TITLE = (
    "CrossCity-TrafficFM: A Reproducible Benchmark Protocol for "
    "Cross-City Mobility Forecasting and Foundation-Model Readiness"
)
AFFILIATION = "National School of Applied Sciences of Tetouan, Abdelmalek Essaadi University, Morocco"
DATE_TEXT = "June 24, 2026"

TABLE_CAPTIONS = {
    "table_01_dataset_summary": "Dataset summary after source aggregation and lag construction.",
    "table_02_model_metrics": "Overall test metrics by model.",
    "table_03_model_metrics_by_city": "City-level test metrics by model.",
    "table_04_ablation_study": "Ablation-study metrics.",
    "table_05_robustness_stress_days": "Stress-day diagnostic metrics.",
    "table_06_morocco_case_study": "Morocco OSM readiness proxy by city.",
    "table_07_claim_evidence_ledger": "Claim and evidence ledger.",
    "table_08_readiness_audit": "Submission-readiness audit.",
    "table_09_paired_tests": "Paired tests against persistence.",
    "table_10_normalized_metrics": "Scale-aware normalized metrics.",
    "table_11_leave_one_city_out": "Leave-one-city-out transfer diagnostic.",
}


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in [
        ("Heading 1", 15, RGBColor(0x12, 0x3B, 0x63)),
        ("Heading 2", 12.5, RGBColor(0x16, 0x55, 0x7A)),
        ("Heading 3", 11.5, RGBColor(0x30, 0x4A, 0x5E)),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)


def clean_inline(text: str) -> str:
    text = text.replace("\\textemdash{}", "\u2014")
    text = text.replace("\\'e", "e")
    text = text.replace("--", "\u2013")
    text = text.replace("\\%", "%")
    text = text.replace("\\_", "_")
    text = text.replace("\\&", "&")
    text = text.replace("~", " ")
    text = re.sub(r"\\texttt\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\textbf\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\small\s*", "", text)
    text = re.sub(r"\\cite\{([^{}]*)\}", lambda m: "[" + m.group(1).replace(",", "; ") + "]", text)
    text = re.sub(r"\\ref\{[^{}]*\}", "", text)
    text = text.replace("Table  .", "Table.")
    text = text.replace("Figure  .", "Figure.")
    text = re.sub(r"\$R\^2\$", "R2", text)
    text = re.sub(r"\$([^$]+)\$", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?\{([^{}]*)\}", r"\1", text)
    text = text.replace("\\", "")
    return re.sub(r"\s+", " ", text).strip()


def extract_environment(tex: str, name: str) -> str:
    m = re.search(rf"\\begin\{{{name}\}}(.*?)\\end\{{{name}\}}", tex, re.S)
    return m.group(1).strip() if m else ""


def extract_keywords(tex: str) -> str:
    m = re.search(r"\\textbf\{Keywords:\}\s*&?\s*(.*?)(?:\\end\{tabular\}|\\end\{minipage\}|\\vspace|\])", tex, re.S)
    return clean_inline(m.group(1)) if m else ""


def add_title_block(doc: Document, tex: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

    author_table = doc.add_table(rows=3, cols=2)
    author_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    author_table.autofit = False
    author_table.columns[0].width = Inches(3.0)
    author_table.columns[1].width = Inches(3.0)
    author_table.cell(0, 0).text = "Ismail Zrigui"
    author_table.cell(0, 1).text = "Samira Khoulji"
    author_table.cell(1, 0).text = AFFILIATION
    author_table.cell(1, 1).text = AFFILIATION
    author_table.cell(2, 0).text = "izrigui@uae.ac.ma"
    author_table.cell(2, 1).text = "skhoulji@uae.ac.ma"
    for row_index, row in enumerate(author_table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(11 if row_index == 0 else 9)

    date_para = doc.add_paragraph(DATE_TEXT)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    abstract = clean_inline(extract_environment(tex, "abstract"))
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(abstract)

    keywords = extract_keywords(tex)
    if keywords:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.78)
        p.paragraph_format.first_line_indent = Inches(-0.78)
        label = p.add_run("Keywords: ")
        label.bold = True
        label.font.name = "Calibri"
        label.font.size = Pt(10.5)
        body = p.add_run(keywords)
        body.font.name = "Calibri"
        body.font.size = Pt(10.5)


def add_csv_table(doc: Document, table_id: str) -> None:
    csv_path = ROOT / "results" / "tables" / f"{table_id}.csv"
    if not csv_path.exists():
        doc.add_paragraph(f"[Table unavailable: {table_id}]")
        return
    caption = TABLE_CAPTIONS.get(table_id, table_id.replace("_", " "))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{table_id}: {caption}")
    run.bold = True

    with csv_path.open(newline="", encoding="utf-8") as handle:
        rows = list(csv.reader(handle))
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(8)
                    if i == 0:
                        run.bold = True


def add_figure(doc: Document, block: str) -> None:
    image_match = re.search(r"\\includegraphics(?:\[[^\]]*\])?\{([^{}]+)\}", block)
    caption_match = re.search(r"\\caption\{(.*?)\}", block, re.S)
    if not image_match:
        return
    stem = Path(image_match.group(1)).stem
    png_path = ROOT / "results" / "figures" / f"{stem}.png"
    if png_path.exists():
        doc.add_picture(str(png_path), width=Inches(5.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = clean_inline(caption_match.group(1)) if caption_match else stem
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True


def add_paragraph_or_list(doc: Document, paragraph: str) -> None:
    paragraph = paragraph.strip()
    if not paragraph:
        return
    if paragraph.startswith("ITEM:"):
        doc.add_paragraph(clean_inline(paragraph[5:]), style="List Number")
        return
    if paragraph.startswith("MATH:"):
        p = doc.add_paragraph(clean_inline(paragraph[5:]))
        for run in p.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        return
    doc.add_paragraph(clean_inline(paragraph))


def split_body_items(body: str) -> list[tuple[str, str]]:
    body = re.sub(r"\\begin\{enumerate\}", "", body)
    body = re.sub(r"\\end\{enumerate\}", "", body)
    body = re.sub(r"\\item\s+", "\n\nITEM:", body)
    body = re.sub(r"\\\[(.*?)\\\]", lambda m: "\n\nMATH:" + m.group(1) + "\n\n", body, flags=re.S)

    tokens: list[tuple[str, str]] = []
    pos = 0
    pattern = re.compile(
        r"(\\section\*?\{[^{}]+\}|\\subsection\{[^{}]+\}|\\subsubsection\{[^{}]+\}|"
        r"\\input\{\.\./results/tables/([^{}]+)\.tex\}|"
        r"\\begin\{figure\*?\}.*?\\end\{figure\*?\})",
        re.S,
    )
    for match in pattern.finditer(body):
        prior = body[pos : match.start()]
        for paragraph in re.split(r"\n\s*\n", prior):
            if paragraph.strip():
                tokens.append(("paragraph", paragraph.strip()))
        text = match.group(0)
        if text.startswith("\\section"):
            tokens.append(("h1", re.search(r"\{([^{}]+)\}", text).group(1)))
        elif text.startswith("\\subsection"):
            tokens.append(("h2", re.search(r"\{([^{}]+)\}", text).group(1)))
        elif text.startswith("\\subsubsection"):
            tokens.append(("h3", re.search(r"\{([^{}]+)\}", text).group(1)))
        elif text.startswith("\\input"):
            tokens.append(("table", Path(match.group(2)).stem))
        elif text.startswith("\\begin{figure"):
            tokens.append(("figure", text))
        pos = match.end()
    for paragraph in re.split(r"\n\s*\n", body[pos:]):
        if paragraph.strip():
            tokens.append(("paragraph", paragraph.strip()))
    return tokens


def add_bibliography(doc: Document) -> None:
    doc.add_heading("References", level=1)
    text = BIB.read_text(encoding="utf-8", errors="replace")
    entries = re.findall(r"@\w+\{([^,]+),(.*?)(?=\n@\w+\{|\Z)", text, flags=re.S)
    for i, (key, fields) in enumerate(entries, start=1):
        title = re.search(r"title\s*=\s*\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}", fields, re.I | re.S)
        author = re.search(r"author\s*=\s*\{([^{}]*)\}", fields, re.I | re.S)
        year = re.search(r"year\s*=\s*\{([^{}]*)\}", fields, re.I | re.S)
        doi = re.search(r"doi\s*=\s*\{([^{}]*)\}", fields, re.I | re.S)
        url = re.search(r"url\s*=\s*\{([^{}]*)\}", fields, re.I | re.S)
        pieces = [f"[{i}]"]
        if author:
            pieces.append(clean_inline(author.group(1).replace(" and ", "; ")))
        if title:
            pieces.append(clean_inline(title.group(1)))
        if year:
            pieces.append(clean_inline(year.group(1)))
        if doi:
            pieces.append("DOI: " + clean_inline(doi.group(1)))
        elif url:
            pieces.append(clean_inline(url.group(1)))
        doc.add_paragraph(" ".join(pieces))


def main() -> None:
    tex = TEX.read_text(encoding="utf-8", errors="replace")
    doc = Document()
    configure_styles(doc)
    add_title_block(doc, tex)

    body_match = re.search(r"\\section\{Introduction\}(.*?)(?:\\FloatBarrier|\\bibliographystyle)", tex, re.S)
    if not body_match:
        raise RuntimeError("Could not locate manuscript body in paper/main.tex")
    body = "\\section{Introduction}" + body_match.group(1)
    for kind, value in split_body_items(body):
        if kind == "h1":
            doc.add_heading(clean_inline(value), level=1)
        elif kind == "h2":
            doc.add_heading(clean_inline(value), level=2)
        elif kind == "h3":
            doc.add_heading(clean_inline(value), level=3)
        elif kind == "table":
            add_csv_table(doc, value)
        elif kind == "figure":
            add_figure(doc, value)
        else:
            add_paragraph_or_list(doc, value)

    add_bibliography(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = TITLE
    doc.core_properties.author = "Ismail Zrigui; Samira Khoulji"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
